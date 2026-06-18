import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def create_report():
    doc = Document()

    # Base styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # 1. Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SPEECH & NLP MODELS BENCHMARK EVALUATION REPORT")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)  # Navy Blue

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("Boilerplate Setup, STT & TTS Performance Analysis")
    run_sub.font.size = Pt(14)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(127, 127, 127)

    doc.add_paragraph("\n")

    # 2. Executive Summary & Recommendation
    h1 = doc.add_heading(level=1)
    run = h1.add_run("1. Executive Summary & GO/NO-GO Recommendation")
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph(
        "This report summarizes the verification, integration, and performance benchmarks of "
        "local Speech-to-Text (STT), Text-to-Speech (TTS), and Natural Language Processing (NLP) "
        "models evaluated on local development hardware (macOS ARM64 CPU). The primary goal is to establish "
        "base performance indicators and determine the production viability of local execution versus cloud services."
    )

    h2 = doc.add_heading(level=2)
    run = h2.add_run("1.1 Recommendation Summary")
    run.font.color.rgb = RGBColor(46, 116, 181)

    # STT Recommendation
    doc.add_heading("Speech-to-Text (STT) - Faster-Whisper: GO (Base Model)", level=3)
    doc.add_paragraph(
        "Based on evaluations across 11 audio clips, the Whisper Base model is recommended. "
        "It achieves a Word Error Rate (WER) of 18.68% while maintaining a low processing latency "
        "of 7.85 seconds per minute of audio. Whisper Tiny is faster (5.04s) but shows an increased error rate (20.74%). "
        "Whisper Small is slower (21.24s) and did not show significant accuracy improvements on local CPU execution."
    )

    # TTS Recommendation
    doc.add_heading("Text-to-Speech (TTS) - Piper: GO (Amy or Joe Voices)", level=3)
    doc.add_paragraph(
        "All 5 evaluated Piper voices demonstrated excellent speed, executing roughly 5x faster than real-time. "
        "Amy (Medium) and Joe (Medium) are the top candidates. Amy achieved the lowest average Real-Time Factor (RTF) of 0.189 "
        "on product phrases, providing high-quality natural voice synthesis under minimal CPU load."
    )

    # NLP Recommendation
    doc.add_heading("Natural Language Processing (NLP) - spaCy: GO", level=3)
    doc.add_paragraph(
        "spaCy's en_core_web_sm pipeline successfully completed verification for entity extraction and linguistic analysis. "
        "It loads instantaneously and runs with sub-millisecond execution times, making it a definitive GO for integration."
    )

    doc.add_page_break()

    # 3. Project Architecture
    h1 = doc.add_heading(level=1)
    run = h1.add_run("2. Repository Architecture & Git Setup")
    run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph(
        "The project layout has been initialized to segregate model runtime services, datasets, and testing configurations. "
        "Local branch protection has been established using Git pre-push hooks to secure the 'main' branch."
    )

    # Folders list
    doc.add_paragraph("Core Directory Segregations:")
    doc.add_paragraph("- backend/: Main backend service architecture.", style='List Bullet')
    doc.add_paragraph("- frontend/: User interface client files.", style='List Bullet')
    doc.add_paragraph("- services/: Shared pipeline loaders (nlp, stt, tts).", style='List Bullet')
    doc.add_paragraph("- data/: Local evaluation datasets (audio/ & transcripts/).", style='List Bullet')
    doc.add_paragraph("- tests/: Evaluation, validation, and benchmarking scripts.", style='List Bullet')
    doc.add_paragraph("- docs/: Generated reports, logs, and documentation.", style='List Bullet')

    doc.add_paragraph("\n")

    # 4. Whisper STT Benchmark Table
    h1 = doc.add_heading(level=1)
    run = h1.add_run("3. Speech-to-Text (STT) Performance Benchmarks")
    run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph(
        "The Whisper STT models (Tiny, Base, and Small) were benchmarked using 11 audio clips (custom recordings + "
        "LJ Speech datasets). Latency represents processing CPU seconds per minute of audio."
    )

    stt_table = doc.add_table(rows=4, cols=3)
    stt_table.style = 'Light Shading Accent 1'
    hdr_cells = stt_table.rows[0].cells
    hdr_cells[0].text = 'Model Size'
    hdr_cells[1].text = 'Latency (Processing Time per 1 Min Audio)'
    hdr_cells[2].text = 'Word Error Rate (WER %)'

    # Shading the header cells
    for cell in hdr_cells:
        set_cell_background(cell, '1F4E79')
        for p_cell in cell.paragraphs:
            for run_cell in p_cell.runs:
                run_cell.font.bold = True
                run_cell.font.color.rgb = RGBColor(255, 255, 255)

    data_stt = [
        ("Tiny", "5.04 seconds", "20.74%"),
        ("Base", "7.85 seconds", "18.68%"),
        ("Small", "21.24 seconds", "17.33%")
    ]

    for idx, (m, lat, wer) in enumerate(data_stt):
        row_cells = stt_table.rows[idx+1].cells
        row_cells[0].text = m
        row_cells[1].text = lat
        row_cells[2].text = wer
        if idx % 2 == 1:
            for cell in row_cells:
                set_cell_background(cell, 'F2F4F7')

    doc.add_paragraph("\n")

    # 5. Piper TTS Benchmark Table
    h1 = doc.add_heading(level=1)
    run = h1.add_run("4. Text-to-Speech (TTS) Voice Benchmarks")
    run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph(
        "5 English voices were evaluated using 5 real product text phrases. "
        "Average Real-Time Factor (RTF) measures the ratio of synthesis CPU time to audio duration."
    )

    tts_table = doc.add_table(rows=6, cols=4)
    tts_table.style = 'Light Shading Accent 1'
    hdr_tts = tts_table.rows[0].cells
    hdr_tts[0].text = 'Voice Model'
    hdr_tts[1].text = 'Total Audio Duration'
    hdr_tts[2].text = 'Total CPU Synthesis Time'
    hdr_tts[3].text = 'Average RTF'

    # Shading the header cells
    for cell in hdr_tts:
        set_cell_background(cell, '1F4E79')
        for p_cell in cell.paragraphs:
            for run_cell in p_cell.runs:
                run_cell.font.bold = True
                run_cell.font.color.rgb = RGBColor(255, 255, 255)

    data_tts = [
        ("Amy (Medium)", "22.07s", "4.07s", "0.189"),
        ("Joe (Medium)", "19.05s", "3.78s", "0.204"),
        ("Danny (Low)", "17.63s", "3.57s", "0.210"),
        ("Lessac (Medium)", "18.03s", "3.71s", "0.211"),
        ("Ryan (Medium)", "17.82s", "3.79s", "0.222")
    ]

    for idx, (voice, dur, cpu, rtf) in enumerate(data_tts):
        row_cells = tts_table.rows[idx+1].cells
        row_cells[0].text = voice
        row_cells[1].text = dur
        row_cells[2].text = cpu
        row_cells[3].text = rtf
        if idx % 2 == 1:
            for cell in row_cells:
                set_cell_background(cell, 'F2F4F7')

    doc.add_paragraph("\n")

    # 6. Conclusion
    h1 = doc.add_heading(level=1)
    run = h1.add_run("5. Project Next Steps")
    run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph("The immediate milestones following this successful evaluation phase include:")
    doc.add_paragraph("1. Integrate verified loaders (spaCy, Faster-Whisper, Piper) into main services/ layers.", style='List Number')
    doc.add_paragraph("2. Develop the backend skeleton REST endpoints inside backend/ to expose STT, TTS, and NLP engines.", style='List Number')
    doc.add_paragraph("3. Establish cloud integration comparison tests (Deepgram STT & ElevenLabs TTS) to evaluate voice quality side-by-side with local hosting.", style='List Number')

    # Save document
    docs_dir = "docs"
    os.makedirs(docs_dir, exist_ok=True)
    report_path = os.path.join(docs_dir, "benchmark_report.docx")
    doc.save(report_path)
    print(f"Report document successfully saved to: {report_path}")

if __name__ == "__main__":
    create_report()
